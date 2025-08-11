import os
import sys
import subprocess
import zipfile
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

# ===== Ensure python-docx is Installed =====
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 0. Kaggle Titanic Dataset Setup =====
if not os.path.exists("train.csv"):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "kaggle"])
        subprocess.check_call([
            "kaggle", "competitions", "download", "-c", "titanic", "-f", "train.csv.zip", "-p", "."
        ])
        with zipfile.ZipFile("train.csv.zip", 'r') as zip_ref:
            zip_ref.extractall(".")
    except Exception as e:
        print("❌ Kaggle download failed. Ensure kaggle.json API key is set up.")
        sys.exit(1)

# ===== 1. Load dataset =====
df = pd.read_csv("train.csv")

# ===== 2. Handle missing numeric values for safe calculations =====
numeric_cols = ["Age", "SibSp", "Parch", "Fare"]
df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].median())

# ===== Descriptive Survival Analysis =====
survival_rate = df['Survived'].mean() * 100
survival_by_gender = df.groupby('Sex')['Survived'].mean() * 100
survival_by_class = df.groupby('Pclass')['Survived'].mean() * 100
avg_age_surv = df.groupby('Survived')['Age'].mean()
embarked_counts = df['Embarked'].value_counts()
embarked_survival = df.groupby('Embarked')['Survived'].mean() * 100
fare_by_class = df.groupby('Pclass')['Fare'].mean()

# ===== Numeric Analysis =====
central_tendency = pd.DataFrame({
    'Mean': df[numeric_cols].mean(),
    'Median': df[numeric_cols].median(),
    'Mode': df[numeric_cols].mode().iloc[0]
})
dispersion = pd.DataFrame({
    'Range': df[numeric_cols].max() - df[numeric_cols].min(),
    'Variance': df[numeric_cols].var(),
    'Std Dev': df[numeric_cols].std(),
    'IQR': df[numeric_cols].quantile(0.75) - df[numeric_cols].quantile(0.25)
})

# ===== Generate Plots =====
sns.set(style="whitegrid")

plt.figure(figsize=(10, 6))
df[numeric_cols].boxplot()
plt.title("Boxplots of Titanic Numeric Columns")
plt.savefig("boxplots_titanic.png")
plt.close()

fig, axes = plt.subplots(2, 2, figsize=(12, 8))
for ax, col in zip(axes.flatten(), numeric_cols):
    sns.histplot(df[col], kde=True, ax=ax, bins=30)
    ax.set_title(f'Histogram of {col}')
plt.tight_layout()
plt.savefig("histograms_titanic.png")
plt.close()

# ===== Create Word Report =====
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_heading("Titanic Data Analysis Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===== Descriptive Survival Analysis Section =====
doc.add_heading("Descriptive Survival Analysis", level=1)
doc.add_paragraph(f"Overall survival rate: {survival_rate:.2f}%")

# Gender table
doc.add_heading("Survival Rate by Gender", level=2)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Gender"
table.rows[0].cells[1].text = "Survival Rate (%)"
for gender, rate in survival_by_gender.items():
    row = table.add_row().cells
    row[0].text = gender.capitalize()
    row[1].text = f"{rate:.2f}"

# Class table
doc.add_heading("Survival Rate by Class", level=2)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Class"
table.rows[0].cells[1].text = "Survival Rate (%)"
for pclass, rate in survival_by_class.items():
    row = table.add_row().cells
    row[0].text = str(pclass)
    row[1].text = f"{rate:.2f}"

# Average age
doc.add_heading("Average Age: Survivors vs Non-Survivors", level=2)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Survived"
table.rows[0].cells[1].text = "Average Age"
for surv, age in avg_age_surv.items():
    row = table.add_row().cells
    row[0].text = "Yes" if surv == 1 else "No"
    row[1].text = f"{age:.2f}"

# Embarked stats
doc.add_heading("Passengers by Embarkation Port", level=2)
table = doc.add_table(rows=1, cols=3)
table.rows[0].cells[0].text = "Port"
table.rows[0].cells[1].text = "Passenger Count"
table.rows[0].cells[2].text = "Survival Rate (%)"
for port, count in embarked_counts.items():
    row = table.add_row().cells
    row[0].text = port
    row[1].text = str(count)
    row[2].text = f"{embarked_survival[port]:.2f}"

# Fare by class
doc.add_heading("Average Fare by Class", level=2)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Class"
table.rows[0].cells[1].text = "Average Fare"
for pclass, fare in fare_by_class.items():
    row = table.add_row().cells
    row[0].text = str(pclass)
    row[1].text = f"{fare:.2f}"

doc.add_page_break()

# ===== Statistical Analysis =====
doc.add_heading("Central Tendency Measures", level=1)
table = doc.add_table(rows=1, cols=4)
hdr = table.rows[0].cells
hdr[0].text = "Column"
hdr[1].text = "Mean"
hdr[2].text = "Median"
hdr[3].text = "Mode"
for col in central_tendency.index:
    row = table.add_row().cells
    row[0].text = col
    row[1].text = f"{central_tendency.loc[col,'Mean']:.2f}"
    row[2].text = f"{central_tendency.loc[col,'Median']:.2f}"
    row[3].text = f"{central_tendency.loc[col,'Mode']:.2f}"

doc.add_heading("Dispersion Measures", level=1)
table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
hdr[0].text = "Column"
hdr[1].text = "Range"
hdr[2].text = "Variance"
hdr[3].text = "Std Dev"
hdr[4].text = "IQR"
for col in dispersion.index:
    row = table.add_row().cells
    row[0].text = col
    row[1].text = f"{dispersion.loc[col,'Range']:.2f}"
    row[2].text = f"{dispersion.loc[col,'Variance']:.2f}"
    row[3].text = f"{dispersion.loc[col,'Std Dev']:.2f}"
    row[4].text = f"{dispersion.loc[col,'IQR']:.2f}"

# Plots
doc.add_heading("Boxplots", level=1)
doc.add_picture("boxplots_titanic.png", width=Inches(5.5))
doc.add_heading("Histograms", level=1)
doc.add_picture("histograms_titanic.png", width=Inches(5.5))

# Conclusion
doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "The Titanic dataset reveals survival patterns linked to gender, class, and ticket price. "
    "Numeric features like 'Fare' and 'Age' show skewness and outliers, requiring preprocessing "
    "for accurate modeling."
)

doc.save("Titanic_Analysis_Report.docx")
print("✅ Report generated: Titanic_Analysis_Report.docx")
